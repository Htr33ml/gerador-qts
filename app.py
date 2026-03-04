from flask import Flask, request, send_file, render_template
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Table, TableStyle
import os

app = Flask(__name__)

# ==========================================================
# FUNÇÕES AUXILIARES
# ==========================================================

def set_font(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

def merge_vertical(table, col, start_row, end_row):
    cell_start = table.cell(start_row, col)
    for row in range(start_row + 1, end_row + 1):
        cell_start.merge(table.cell(row, col))

def merge_horizontal(row, start_col, end_col):
    cell = row.cells[start_col]
    for col in range(start_col + 1, end_col + 1):
        cell.merge(row.cells[col])
    return cell

# ==========================================================
# ROTA PRINCIPAL
# ==========================================================

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":

        semana = request.form.get("semana")
        periodo = request.form.get("periodo")
        companhia = request.form.get("companhia")

        dias_sem_expediente = request.form.getlist("sem_expediente")

        # Criar documento base
        doc = Document("modelo_qts.docx")

        # Substituir dados fixos
        for p in doc.paragraphs:
            if "{{SEMANA}}" in p.text:
                p.text = p.text.replace("{{SEMANA}}", semana)
            if "{{PERIODO}}" in p.text:
                p.text = p.text.replace("{{PERIODO}}", periodo)
            if "{{COMPANHIA}}" in p.text:
                p.text = p.text.replace("{{COMPANHIA}}", companhia)

        # Encontrar marcador {{DIAS_SEMANA}}
        for i, p in enumerate(doc.paragraphs):
            if "{{DIAS_SEMANA}}" in p.text:
                insert_index = i
                p.text = ""
                break

        dias = ["SEGUNDA", "TERÇA", "QUARTA", "QUINTA", "SEXTA"]

        for dia in dias:

            doc.paragraphs[insert_index].insert_paragraph_before("")
            doc.paragraphs[insert_index].insert_paragraph_before("")

            # Criar tabela
            table = doc.add_table(rows=1, cols=9)
            headers = ["DATA", "HORA", "INSTRUENDO", "LOCAL", "UNIFORME",
                       "INSTRUTOR", "MATÉRIA", "OII", "OBJETIVO"]

            for col, header in enumerate(headers):
                table.rows[0].cells[col].text = header
                set_font(table.rows[0].cells[col])

            if dia in dias_sem_expediente:
                row = table.add_row()
                cell = merge_horizontal(row, 0, 8)
                cell.text = "SEM EXPEDIENTE"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(cell)

            else:
                atividades = request.form.getlist(f"atividade_{dia}")

                start_row = 1

                for atividade in atividades:
                    dados = atividade.split("|")
                    row = table.add_row()

                    # DATA será preenchida depois
                    row.cells[1].text = dados[0]  # hora
                    set_font(row.cells[1])

                    if dados[1] == "geral":
                        cell = merge_horizontal(row, 2, 8)
                        cell.text = dados[2]
                        set_font(cell)
                    else:
                        for i in range(2, 9):
                            row.cells[i].text = dados[i-1]
                            set_font(row.cells[i])

                # Escrever DATA na primeira linha
                if len(table.rows) > 1:
                    table.rows[1].cells[0].text = f"{dia}"
                    set_font(table.rows[1].cells[0])
                    merge_vertical(table, 0, 1, len(table.rows)-1)

        doc.save("QTS.docx")

        # ==================================================
        # GERAR PDF
        # ==================================================

        pdf_file = "QTS.pdf"
        pdf = SimpleDocTemplate(pdf_file, pagesize=A4)
        elements = []

        style = ParagraphStyle(
            name="Normal",
            fontName="Times-Roman",
            fontSize=11
        )

        elements.append(Paragraph("QTS gerado automaticamente.", style))
        elements.append(Spacer(1, 12))

        pdf.build(elements)

        return send_file("QTS.docx", as_attachment=True)

    return render_template("index.html")

# ==========================================================

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
    
